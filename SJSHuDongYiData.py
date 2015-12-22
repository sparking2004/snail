
# -*- coding: cp936 -*-
import mechanize
from bs4 import BeautifulSoup
import time
import re
import urllib2
import os

#doc
from docx import Document
from docx.shared import Inches
from win32com import client

#pdf
from pdfminer.layout import LTTextContainer
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice

#br = mechanize.Browser()
#response = br.open('http://irm.cninfo.com.cn/ircs/interaction/irmInformationList.do?pageNo=1&stkcode=&beginDate=2015-10-30&endDate=2015-11-30&keyStr=&irmType=251314')
#print response.read()



#soup = BeautifulSoup(open(r'C:\Users\zcj\Desktop\text2.html'))


#from docx import Document
#from docx.shared import Inches
#from PyPDF2 import PdfFileWriter, PdfFileReader
#import pdfminer
#document = Document(r'c:\1.docx')
#tables = document.tables

#for table in tables:
#    for row in table.rows:
#        for cell in row.cells:
#            print(cell.text)
 #           print('++++++++++++++++')


#input1 = PdfFileReader(open(r"c:\1.pdf", "rb"))
#print input1.getDocumentInfo()
#f = open("c:\\2.txt","w")
#input1.read(f)
#print "document1.pdf has %d pages." % input1.getNumPages()
#page = input1.getPage(1)
#print page['/Contents']
#print page.getContents()['/Filter']

def CreateDirs(dirs):
    if not os.path.exists(dirs):
        print u'����Ŀ¼:'+ dirs
        os.makedirs(dirs)

class InvestmentInfo(object):
    def __init__(self):
        self.code = ""
        self.name = ""
        self.uploadDate = ""
        self.recordData = ""

        #�����׼�¼�ļ�����洢·��
        self.recordFileAddr =""

        self.fileExt = ""

        #�����׼�¼���ش洢·��
        self.localAddr = ""

        #���л�������
        self.company = []

class InvestmentInfoTable(object):
    def __init__(self):
        self._addr = r'http://irm.cninfo.com.cn/ircs/interaction/irmInformationList.do?pageNo=1&stkcode=&beginDate=%s&endDate=%s&keyStr=&irmType=251314'

    def GetInvestmentInfo(self,investmentInfoVec,beginDate,endDate):
        '''
        ��ȡ��¼��Ϣ
        :param investmentInfoVec: ��ȡ����Ϣ�ṹ��
        :param beginDate: ��ѯ�Ŀ�ʼʱ�䣬����2015-12-01
        :param endDate: ��ѯ�Ľ���ʱ�䣬����2015-12-01
        :return:
        '''
        soup = BeautifulSoup(self._GetWebPageInfo(beginDate,endDate))
        for tr in soup.table.find_all('tr'):
            tagAs = tr.find_all('a')
            if len(tagAs) == 2:
                print tagAs[0]
                print tagAs[1].decode()
                info = InvestmentInfo()
                self._GetCode(tagAs[0],info)
                self._GetOtherInfo(tagAs[1],info)
                investmentInfoVec.append(info)

            else:
                print 'TagA num error!'

    def GetRecordFile(self,investmentInfoVec,downloadDir):
        '''
        ���ݼ�¼��Ϣ�����ļ��������investmentInfo�ṹ��ı����ļ���Ϣ
        :param investmentInfoVec: investmentInfo��Ϣ�ṹ��
        :param downloadDir: ��Ҫ���ļ����صı���Ŀ¼
        :return:
        '''
        for info in investmentInfoVec:
            filedir = downloadDir + info.uploadDate
            CreateDirs(filedir)
            filename = re.sub(r'[/\\*?<>:|]',"",info.name)
            filename = filedir+'\\'+info.code+filename+info.recordData+'.'+info.fileExt
            f = urllib2.urlopen(info.recordFileAddr)
            with open(filename, "wb") as code:
                code.write(f.read())


    def _GetWebPageInfo(self,beginDate,endDate):
        pageAddr = self._addr % (beginDate,endDate)
        br = mechanize.Browser()
        response = br.open(pageAddr)
        return response.read()

    def _GetCode(self,tagA,info):
        info.code = unicode(tagA.string).strip()

    def _GetOtherInfo(self,tagA,info):
        info.recordFileAddr = tagA['href']
        #������¼���ں��ļ���׺��
        m = re.match('.+/([0-9]{4}-[0-9]{1,2}-[0-9]{1,2})/.+\.(.+)\?.+',info.recordFileAddr)
        if m is None:
            print '��ȡ�ϴ����ڼ���չ������'
        info.uploadDate =m.group(1)
        info.fileExt = m.group(2).lower()
        title = tagA['title']
        m = re.match(u'(.+) *�� *([0-9]{4}��[0-9]{1,2}��[0-9]{1,2}��).*',tagA['title'])
        if m is None:
            print '��ȡ��¼���ڼ����ƴ���'
        info.name = m.group(1)
        info.recordData = m.group(2)


class InvestigateInfo(object):
    def GetInvestigateCompanyAndPeople(self,file):
        '''
        :param file: Ͷ���߻������ĵ�·����֧��pdf��doc�� docx
        :return:���ص��л���������һ���б�
        '''
        companyInfo = ''
        if file.endswith(r'.doc'):
            companyInfo = self._GetFromDoc(file)
        elif file.endswith(r'.docx'):
            companyInfo = self._GetFromDocx(file)
        elif file.endswith(r'.pdf'):
            companyInfo = self._GetFromPdf(file)

        return self._FormatCompanyInfo(companyInfo)


    def _FormatCompanyInfo(self,compInfo):
        compInfo = compInfo.replace(u'��',r'\r')
        return [x.strip() for x in compInfo.split(r'\r') if len(x.strip())!=0]


    def _GetFromPdf(self,pdf):
        '''
        �ο��ĵ�http://www.unixuser.org/~euske/python/pdfminer/programming.html
        '''
        pass
        fp = open(pdf, 'rb')
        #���ļ�����������һ��pdf�ĵ�������
        parser = PDFParser(fp)
        # ����һ��  PDF �ĵ�
        doc = PDFDocument(parser)
        # ���ӷ����� ���ĵ�����
        parser.set_document(doc)
        # ����ĵ��Ƿ��ṩtxtת�������ṩ�ͺ���
        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed

        # ����PDf ��Դ������ ����������Դ
        rsrcmgr = PDFResourceManager()
        # ����һ��PDF�豸����
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
            # receive the LTPage object for the page.
            layout = device.get_result()
            for x in layout:
                if(isinstance(x, LTTextContainer)):
                    print x.get_text()

        pass

    def _GetFromDoc(self,wordDoc):
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(wordDoc)
        return doc.Tables[0].Rows[1].Cells[1].Range.Text

    def _GetFromDocx(self,wordDocx):
        document = Document(wordDocx)
        return document.tables[0].cell(1,1).text




if __name__ == "__main__":

    #curTime = '2015-12-21'
    curTime = time.strftime('%Y-%m-%d', time.localtime(time.time()))
    web = InvestmentInfoTable()
    infoVec = []
    web.GetInvestmentInfo(infoVec,curTime,curTime)
    #print infoVec

    #web.GetRecordFile(infoVec,'D:\\test\\')

    #��ӡcode���ı�
    f = open('D:\\test\\1.txt','w')
    codevec = []
    for info in infoVec:
        codevec.append(info.code)
    f.write('\n'.join(codevec))
    f.close()


    read = InvestigateInfo()
    s = read.GetInvestigateCompanyAndPeople(u'D:\\test\\2.pdf')
    for li in s:
        print li






